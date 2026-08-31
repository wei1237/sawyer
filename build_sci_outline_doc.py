from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from pathlib import Path


OUT = Path(r"D:\ubuntu20\SCI机器人论文大纲初稿_20260823.docx")


def set_run_font(run, size=None, bold=None, color=None, font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_para_font(paragraph, size=None, bold=None, color=None, font="微软雅黑"):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color, font=font)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False, fill=None):
    cell.text = ""
    p = cell.paragraphs[0]
    r = p.add_run(text)
    set_run_font(r, size=9.5, bold=bold, color="000000")
    p.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if fill:
        set_cell_shading(cell, fill)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    r = p.add_run(text)
    set_run_font(r, size=10.5)
    p.paragraph_format.space_after = Pt(3)
    return p


def add_note(doc, title, text):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=10.5, bold=True, color="1F4D78")
    p2 = cell.add_paragraph()
    r2 = p2.add_run(text)
    set_run_font(r2, size=10.5)
    p2.paragraph_format.space_after = Pt(0)
    return table


def add_section(doc, title, purpose, write_points=None, figures=None, caution=None):
    p = doc.add_heading(title, level=1)
    set_para_font(p, size=16, bold=True, color="2E74B5")
    if purpose:
        add_note(doc, "本节作用", purpose)
    if write_points:
        p2 = doc.add_heading("建议写法", level=2)
        set_para_font(p2, size=13, bold=True, color="2E74B5")
        for item in write_points:
            add_bullet(doc, item)
    if figures:
        p3 = doc.add_heading("建议图表/结果位置", level=2)
        set_para_font(p3, size=13, bold=True, color="2E74B5")
        for item in figures:
            add_bullet(doc, item)
    if caution:
        add_note(doc, "写作注意", caution)


def add_subsection(doc, title, bullets):
    p = doc.add_heading(title, level=2)
    set_para_font(p, size=13, bold=True, color="2E74B5")
    for b in bullets:
        add_bullet(doc, b)


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "微软雅黑"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name in ["Heading 1", "Heading 2", "Heading 3", "List Bullet", "List Bullet 2"]:
        st = styles[style_name]
        st.font.name = "微软雅黑"
        st._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于少量示教快速迁移的机器人操作泛化方法")
    set_run_font(r, size=20, bold=True, color="0B2545")
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("SCI 机器人论文大纲初稿（中文标题 + 英文括号版）")
    set_run_font(r, size=12, bold=False, color="555555")
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = meta.add_run("版本：2026-08-23｜用途：论文中文初稿扩写前的结构稿")
    set_run_font(r, size=9.5, color="555555")

    add_note(
        doc,
        "定位说明",
        "本文档按 SCI 机器人论文投稿结构组织，不按毕业论文结构展开。当前版本重点是把“问题—方法—实验—结果—讨论”的逻辑搭好，后续可直接在每个小节下补公式、图表、实验数据和参考文献。"
    )

    p = doc.add_heading("章节功能总览（Paper Structure Overview）", level=1)
    set_para_font(p, size=16, bold=True, color="2E74B5")
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["章节", "核心作用", "需要回答的问题", "建议输出"]
    for i, h in enumerate(headers):
        set_cell_text(table.rows[0].cells[i], h, bold=True, fill="E8EEF5")
    overview_rows = [
        ("1 引言", "提出问题和贡献", "为什么少示教操作泛化值得做？", "问题描述、局限、贡献"),
        ("2 相关工作", "定位差异", "已有学习/迁移/几何方法哪里不够？", "分类综述和差异总结"),
        ("3 问题定义", "形式化任务", "demo、live scene 和迁移目标如何表示？", "符号、变量、目标定义"),
        ("4 方法", "核心技术", "如何从语言、视觉、几何到轨迹执行？", "框架图、公式、模块说明"),
        ("5 实验设置", "说明实验可信度", "任务、平台、指标和协议是什么？", "实验表、评价指标"),
        ("6 仿真结果", "主结果", "不同泛化条件下成功率和误差如何？", "热力图、柱状图、误差图"),
        ("7 真实机器人验证", "验证可迁移性", "仿真结果能否在真实系统复现？", "真实机器人案例图"),
        ("8 结论", "收束贡献", "方法价值、局限和未来方向是什么？", "总结和未来工作"),
    ]
    for row in overview_rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    add_section(
        doc,
        "1 引言（Introduction）",
        "引言要从机器人操作泛化的现实困难出发，逐步引出少量示教、语言检索、视觉感知和局部几何对齐的必要性，最后明确本文贡献。",
        [
            "1.1 研究背景与动机（Background and Motivation）：说明机器人在新位置、新姿态、新形状、新物体关系下执行操作时，重新采集示教或重新训练模型成本高。",
            "1.2 现有方法的局限（Limitations of Existing Approaches）：按学习方法、直接轨迹复现、几何迁移三类写，指出它们分别存在数据量大、坐标系固定、任务级物体关系利用不足的问题。",
            "1.3 基于物体相对几何的示范迁移思想（Object-Relative Geometry-Aware Replay）：强调本文不是重新学习 policy，而是复用已有 demonstration，通过物体相对关系、几何对齐和轨迹迁移实现泛化。",
            "1.4 主要贡献（Contributions）：建议写三点：少示教物体相对几何感知示范复用框架；几何感知局部轨迹适应方法；多任务仿真与真实机器人验证。"
        ],
        [
            "可在引言最后放 Fig. 1 的简化流程图：Task query → Demo retrieval → Semantic perception → Geometric alignment → Trajectory transfer → Execution and logging。",
            "贡献段落建议用编号列出，不要写成泛泛的系统实现。"
        ],
        "不要把方法夸大成完整 3D 点云稠密对应。更稳妥的表述是“任务相关局部几何对齐”。"
    )

    add_section(
        doc,
        "2 相关工作（Related Work）",
        "相关工作要服务于本文定位：少示教、检索式迁移、物体中心/任务关系几何，而不是泛泛介绍所有机器人学习方法。",
        [
            "2.1 面向机器人操作的示范学习（Learning from Demonstration for Robot Manipulation）：介绍 DMP、GMM/GMR、imitation learning、trajectory learning，最后指出本文不是学习新策略，而是复用已有示教轨迹。",
            "2.2 少示范与可泛化机器人操作（Few-Shot and Generalizable Robot Manipulation）：介绍 MT3、retrieval-based manipulation、skill transfer，强调少量 demo 条件下的新场景适应。",
            "2.3 几何表示与任务迁移（Geometric Representation and Task Transfer）：介绍 geometry-aware transfer、object-centric representation、trajectory transformation，突出本文更关注任务/物体几何，而不是单纯 robot-centric geometry。",
            "可补充 SemAnCorr 等 training-free / correspondence 相关工作，但要说明本文不是 dense surface correspondence，而是局部任务几何量驱动的轨迹迁移。"
        ],
        [
            "建议放一个 Related Work 分类表：方法类别、代表方法、依赖数据、几何使用方式、与本文差异。",
        ],
        "相关工作不要写成文献堆砌。每一小节最后都要落到“为什么需要本文方法”。"
    )

    add_section(
        doc,
        "3 问题定义（Problem Formulation）",
        "本章负责把实验系统变成论文中的形式化问题。重点定义示教、当前场景、几何迁移目标和成功评价。",
        [
            "3.1 示范表示（Demonstration Representation）：定义一个 demo 包含任务描述、物体几何、关键瓶颈位姿、操作轨迹和夹爪事件。",
            "3.2 当前场景表示（Live Scene Representation）：说明通过 RGB-D、语义分割和局部点云得到 object pose、object size、top surface、yaw 或物体间相对关系。",
            "3.3 几何迁移目标（Objective of Geometric Transfer）：目标是把 demo trajectory 转换到 live trajectory，同时尽量保持操作点与目标/锚点/插孔之间的相对关系。",
            "可以定义 T_demo、T_live、p_o^demo、p_o^live 等符号，为第 4 章公式铺垫。"
        ],
        [
            "建议放 Table 1：符号定义表。",
            "建议放一个小示意图：demo scene 与 live scene 中物体关系保持一致。"
        ],
        "这里不要塞太多实现细节；具体 LangSAM、MoveIt、Gazebo 放到方法和实验设置里。"
    )

    add_section(
        doc,
        "4 方法（Method）",
        "方法章是论文核心。建议按系统流程写，但要突出两个关键词：object-relative geometry 和 local trajectory adaptation。",
        [
            "4.1 系统总体框架（System Overview）：输入为 task query 和 live RGB-D scene；输出为可执行轨迹和自动记录结果。",
            "4.2 示范库与任务检索（Demonstration Library and Retrieval）：示范库保存 task description、object information、trajectory、gripper action；检索结合 language query 和 geometric similarity。",
            "4.3 语义感知与三维几何估计（Semantic Perception and 3D Geometry Estimation）：LangSAM 生成 mask，结合 depth / PointCloud2 得到局部点云，估计 object center、size、top surface、yaw。",
            "4.4 物体相对几何对齐（Object-Relative Geometric Alignment）：写位置映射和位姿映射公式，例如 p_live = p_o^live + (p_demo - p_o^demo)，以及 T_live = T_o^live (T_o^demo)^-1 T_demo。",
            "4.5 几何感知局部轨迹适应（Geometry-Aware Local Trajectory Adaptation）：作为创新重点展开，包括高度感知映射、抓取口几何锚定修正、局部轨迹融合。",
            "4.6 机器人执行（Robot Execution）：说明 MoveIt Cartesian execution、gripper control、lift/place/insert verification。"
        ],
        [
            "Fig. 1：系统总框图。",
            "Fig. 2：物体相对几何迁移示意图。",
            "Fig. 3：几何感知局部轨迹适应示意图，突出只修正接触前/接触附近轨迹，而不是完全重规划策略。",
            "Algorithm 1：Geometry-Aware Demonstration Replay 的伪代码。"
        ],
        "旋转抓取中 before-close mouth error 可以作为 open-loop replay 的局限分析，不建议在方法章把它包装成所有任务都用了的闭环补偿。"
    )

    add_subsection(doc, "4.5.1 高度感知映射（Height-Aware Mapping）", [
        "解决不同物体高度变化导致的抓取/接触点高度不一致问题，例如 cube demo 迁移到 cylinder。",
        "建议写成：根据目标 top surface 或局部点云高度修正 bottleneck pose 和 grasp/contact pose 的 z 分量。",
        "注意说明这是局部几何修正，不是完整物体重建。"
    ])
    add_subsection(doc, "4.5.2 抓取口几何锚定修正（Mouth Anchor Correction）", [
        "解决系统控制的 right_hand 位姿与真实 gripper mouth center 不完全一致的问题。",
        "普通 top_grasp 可以在闭合前做 mouth-center 微调；rotated_top_grasp 当前策略是记录偏差用于分析，不强行补偿。",
        "论文中建议把该模块写成可选的局部执行修正模块，并在消融实验中说明启用/禁用差异。"
    ])
    add_subsection(doc, "4.5.3 局部轨迹融合（Tail Blending）", [
        "目标是只调整接触前或闭合前局部轨迹，保留原始示教动作风格。",
        "可解释为在原始 replay 轨迹和几何修正目标之间做局部融合，避免全轨迹大幅变形。",
        "如果正式实验中不同任务使用策略不同，需要在实验协议里标注。"
    ])

    add_section(
        doc,
        "5 实验设置（Experimental Setup）",
        "实验设置要让审稿人相信结果可复现、变量清楚、成功判定客观。",
        [
            "5.1 仿真与真实机器人平台（Simulation and Real Robot Platforms）：仿真写 Gazebo、Sawyer、RGB-D camera、MoveIt；真实平台写 Sawyer、ASC60C、LangSAM、camera-base calibration。",
            "5.2 任务与示范设计（Tasks and Demonstrations）：包括位置抓取泛化、姿态抓取泛化、形状泛化抓取、放置任务迁移、插入任务迁移。",
            "5.3 评价指标（Evaluation Metrics）：抓取写 success rate、object retention；放置写 center error、yaw error；插入写 insertion success、alignment error、insert depth。",
            "5.4 实验协议（Experimental Protocol）：说明 trial 数量、重复次数、是否每次重新感知、demo 是否固定、invalid trial 如何处理。"
        ],
        [
            "Table 2：实验任务与泛化变量。",
            "Table 3：评价指标和成功判定阈值。",
            "可补充 failure_category 统计字段：perception failure、planning failure、grasp execution failure、place error、insert collision 等。"
        ],
        "不同任务的泛化变量不必完全一样，但同一个子实验内部变量要清楚。例如有锚点放置看方向影响，插入看初始偏移量影响。"
    )

    p = doc.add_heading("建议任务-变量对应表（可直接改成论文表格）", level=2)
    set_para_font(p, size=13, bold=True, color="2E74B5")
    task_table = doc.add_table(rows=1, cols=4)
    task_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["任务", "主要泛化变量", "建议指标", "建议图"]):
        set_cell_text(task_table.rows[0].cells[i], h, bold=True, fill="E8EEF5")
    for row in [
        ("顶部抓取", "位置、形状/尺寸", "抓取成功率、抬升高度、目标保持", "位置 × 形状热力图"),
        ("旋转顶部抓取", "位置、yaw", "成功率、yaw 误差、before-close mouth error", "位置 × yaw 热力图"),
        ("无锚点放置", "目标位置、方向偏移", "放置中心误差、最终成功率", "方向条件热力图"),
        ("有锚点放置", "锚点位置、目标相对锚点方向", "相对关系误差、成功率", "锚点位置 × 相对方向热力图"),
        ("竖直插入", "圆柱-插孔初始偏移量", "插入成功率、alignment error、insert depth", "偏移量-成功率曲线"),
    ]:
        cells = task_table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], text)

    add_section(
        doc,
        "6 仿真实验结果（Simulation Results）",
        "结果章要先给总体成功率，再按任务展示内部规律，最后用基线和消融说明方法模块的作用。",
        [
            "6.1 位置泛化结果（Position Generalization Results）：展示 workspace 位置变化下的抓取成功率。",
            "6.2 姿态泛化结果（Orientation Generalization Results）：展示 yaw 变化下旋转抓取的成功率和偏抓问题。",
            "6.3 形状泛化结果（Shape Generalization Results）：重点写 cube demo → cylinder 等跨形状泛化，体现 geometry-aware 的必要性。",
            "6.4 放置任务结果（Placement Results）：分别分析无锚点和有锚点条件下相对关系迁移的效果。",
            "6.5 插入任务结果（Insertion Results）：如果测试初始偏移量影响，可画偏移量与成功率/插入误差关系。",
            "6.6 基线比较（Baseline Comparison）：比较 Naive Replay、Translation-only Replay、Geometry-Aware Replay。",
            "6.7 消融实验（Ablation Study）：验证去除几何对齐、去除高度映射、去除抓取口修正后的影响。"
        ],
        [
            "Fig. 4：五类任务总体成功率柱状图。",
            "Fig. 5：五个任务各自热力图。",
            "Fig. 6：失败原因堆叠柱状图。",
            "Fig. 7：感知/几何误差与成功失败关系散点图。",
            "Table 4：基线比较结果。",
            "Table 5：消融实验结果。"
        ],
        "新增字段如果旧数据没有，不必强行重跑所有实验。主结果用所有数据共有字段；新增字段用于失败诊断和补充分析。"
    )

    add_section(
        doc,
        "7 真实机器人验证（Real Robot Validation）",
        "真实机器人部分不一定要覆盖所有仿真条件，重点是展示方法不是只在 Gazebo 里成立。",
        [
            "7.1 真实机器人系统集成（Real Robot System Integration）：ASC60C → LangSAM → Depth Point Cloud → Geometry Transfer → MoveIt → Sawyer。",
            "7.2 真实机器人实验结果（Real Robot Experimental Results）：建议选三个代表任务：位置抓取、跨形状抓取、插入任务。",
            "7.3 仿真到真实分析（Simulation-to-Real Analysis）：讨论 perception noise、calibration error、execution uncertainty 对结果的影响。"
        ],
        [
            "Fig. 8：真实机器人系统照片和坐标系标定示意。",
            "Fig. 9：真实机器人执行序列图，按关键帧展示感知、接近、接触、完成。",
            "Table 6：真实机器人代表任务成功率。"
        ],
        "如果真实机器人实验数量少，表述要写成 representative validation，不要写成大规模真实实验。"
    )

    add_section(
        doc,
        "8 结论（Conclusion）",
        "结论要回到本文解决的问题：少量示教下如何通过语言、视觉和任务相关几何把操作轨迹迁移到新条件。",
        [
            "8.1 工作总结（Summary）：总结提出的框架、几何迁移方法和多任务实验验证。",
            "8.2 局限与未来工作（Limitations and Future Work）：写依赖视觉质量、依赖抓取一致性、交互阶段仍偏开环。",
            "未来方向可写 closed-loop interaction、larger demonstration library、complex objects、real-world long-horizon tasks。"
        ],
        [
            "结论通常不放新图，但可以简洁引用总体成功率和代表性提升。"
        ],
        "不要在结论首次提出新实验或新技术点。"
    )

    p = doc.add_heading("可直接使用的问题描述草稿（Problem Statement Draft）", level=1)
    set_para_font(p, size=16, bold=True, color="2E74B5")
    problem = (
        "机器人在非结构化环境中执行抓取、放置和插入等操作时，常需要面对目标位置、姿态、尺寸、形状以及物体间相对关系的变化。"
        "如果每一种变化都重新采集大量示教或重新训练策略模型，将导致数据采集成本高、任务扩展效率低，并限制机器人在新环境中的快速部署能力。"
        "因此，如何在少量示教条件下，利用语言指令、视觉感知和任务相关几何信息，将已有示教轨迹快速迁移到新的目标状态和任务条件中，是机器人操作泛化中的关键问题。"
        "本文围绕这一问题，研究一种面向少示教机器人操作的物体相对几何感知示范复用框架，并通过多类仿真任务和代表性真实机器人实验评估其成功率、误差和失败模式。"
    )
    p = doc.add_paragraph()
    r = p.add_run(problem)
    set_run_font(r, size=10.5)

    p = doc.add_heading("下一步写作清单（Next Writing Checklist）", level=1)
    set_para_font(p, size=16, bold=True, color="2E74B5")
    checklist = [
        "把每个任务的正式实验条件表固定下来，避免边跑边改导致论文表述混乱。",
        "整理所有 CSV/JSONL 字段，确定主结果使用哪些共同字段，新增诊断字段只做补充分析。",
        "先画总体成功率柱状图和各任务热力图，再决定是否需要额外消融。",
        "方法章中明确区分：完整点云配准、局部几何估计、任务相关几何对齐，避免概念过度包装。",
        "相关工作中补齐 MT3、few-shot manipulation、object-centric representation、training-free transfer/correspondence 方向的文献。",
        "真实机器人部分先做 representative validation，数量不足时不要写成大规模真实实验。"
    ]
    for item in checklist:
        add_bullet(doc, item)

    footer = sec.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = footer.add_run("SCI论文大纲初稿｜基于少量示教快速迁移的机器人操作泛化方法")
    set_run_font(fr, size=8, color="777777")

    doc.save(OUT)


if __name__ == "__main__":
    build()
    print(OUT)
