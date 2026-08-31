from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


OUT = Path(r"D:\ubuntu20\SCI机器人论文目录_可跳转标题正文版_20260823.docx")


def set_font(run, size=11, bold=False, color="000000", font="微软雅黑", underline=False):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.underline = underline
    run.font.color.rgb = RGBColor.from_string(color)


def set_style_font(style, size, bold=False):
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style._element.rPr.rFonts.set(qn("w:ascii"), "微软雅黑")
    style._element.rPr.rFonts.set(qn("w:hAnsi"), "微软雅黑")
    style.font.size = Pt(size)
    style.font.bold = bold
    style.font.color.rgb = RGBColor(0, 0, 0)


def add_bookmark(paragraph, name, bookmark_id):
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def add_internal_hyperlink(paragraph, text, anchor, level):
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")

    run = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "000000")
    rpr.append(color)

    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "none")
    rpr.append(underline)

    size = OxmlElement("w:sz")
    size.set(qn("w:val"), "22" if level <= 2 else "20")
    rpr.append(size)

    size_cs = OxmlElement("w:szCs")
    size_cs.set(qn("w:val"), "22" if level <= 2 else "20")
    rpr.append(size_cs)

    fonts = OxmlElement("w:rFonts")
    fonts.set(qn("w:eastAsia"), "微软雅黑")
    fonts.set(qn("w:ascii"), "微软雅黑")
    fonts.set(qn("w:hAnsi"), "微软雅黑")
    rpr.append(fonts)

    if level == 1:
        bold = OxmlElement("w:b")
        rpr.append(bold)

    run.append(rpr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_toc_entry(doc, text, anchor, level):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.left_indent = Inches(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
    elif level == 2:
        p.paragraph_format.left_indent = Inches(0.28)
        p.paragraph_format.space_after = Pt(2)
    else:
        p.paragraph_format.left_indent = Inches(0.55)
        p.paragraph_format.space_after = Pt(1)
    add_internal_hyperlink(p, text, anchor, level)


def add_body_heading(doc, text, anchor, bookmark_id, level):
    style = f"Heading {level}"
    p = doc.add_paragraph(style=style)
    add_bookmark(p, anchor, bookmark_id)
    run = p.add_run(text)
    set_font(run, size={1: 16, 2: 13, 3: 12}[level], bold=(level == 1), color="000000")
    p.paragraph_format.space_before = Pt({1: 16, 2: 10, 3: 6}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 5, 3: 3}[level])
    return p


def build():
    doc = Document()

    sec = doc.sections[0]
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    set_style_font(doc.styles["Normal"], 11, False)
    set_style_font(doc.styles["Heading 1"], 16, True)
    set_style_font(doc.styles["Heading 2"], 13, False)
    set_style_font(doc.styles["Heading 3"], 12, False)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(14)
    r = title.add_run("基于少量示教快速迁移的机器人操作泛化方法")
    set_font(r, size=18, bold=True, color="000000")

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_title.paragraph_format.space_after = Pt(12)
    r = toc_title.add_run("目录")
    set_font(r, size=15, bold=True, color="000000")

    lines = [
        (1, "1 引言（Introduction）"),
        (2, "1.1 研究背景与动机（Background and Motivation）"),
        (2, "1.2 现有方法的局限（Limitations of Existing Approaches）"),
        (3, "1.2.1 基于学习的方法（Learning-based Manipulation）"),
        (3, "1.2.2 直接轨迹复现方法（Direct Trajectory Replay）"),
        (3, "1.2.3 几何迁移方法（Geometric Transfer Methods）"),
        (2, "1.3 基于物体相对几何的示范迁移思想（Object-Relative Geometry-Aware Replay）"),
        (2, "1.4 主要贡献（Contributions）"),
        (1, "2 相关工作（Related Work）"),
        (2, "2.1 面向机器人操作的示范学习（Learning from Demonstration for Robot Manipulation）"),
        (2, "2.2 少示范与可泛化机器人操作（Few-Shot and Generalizable Robot Manipulation）"),
        (2, "2.3 几何表示与任务迁移（Geometric Representation and Task Transfer）"),
        (1, "3 问题定义（Problem Formulation）"),
        (2, "3.1 示范表示（Demonstration Representation）"),
        (2, "3.2 当前场景表示（Live Scene Representation）"),
        (2, "3.3 几何迁移目标（Objective of Geometric Transfer）"),
        (1, "4 方法（Method）"),
        (2, "4.1 系统总体框架（System Overview）"),
        (2, "4.2 示范库与任务检索（Demonstration Library and Retrieval）"),
        (2, "4.3 语义感知与三维几何估计（Semantic Perception and 3D Geometry Estimation）"),
        (2, "4.4 物体相对几何对齐（Object-Relative Geometric Alignment）"),
        (2, "4.5 几何感知局部轨迹适应（Geometry-Aware Local Trajectory Adaptation）"),
        (3, "4.5.1 高度感知映射（Height-Aware Mapping）"),
        (3, "4.5.2 抓取口几何锚定修正（Mouth Anchor Correction）"),
        (3, "4.5.3 局部轨迹融合（Tail Blending）"),
        (2, "4.6 机器人执行（Robot Execution）"),
        (1, "5 实验设置（Experimental Setup）"),
        (2, "5.1 仿真与真实机器人平台（Simulation and Real Robot Platforms）"),
        (3, "5.1.1 仿真平台（Simulation Platform）"),
        (3, "5.1.2 真实机器人平台（Real Robot Platform）"),
        (2, "5.2 任务与示范设计（Tasks and Demonstrations）"),
        (3, "5.2.1 位置抓取泛化（Position Grasp Generalization）"),
        (3, "5.2.2 姿态抓取泛化（Orientation Grasp Generalization）"),
        (3, "5.2.3 形状泛化抓取（Shape Generalization Grasp）"),
        (3, "5.2.4 放置任务迁移（Placement Generalization）"),
        (3, "5.2.5 插入任务迁移（Insertion Generalization）"),
        (2, "5.3 评价指标（Evaluation Metrics）"),
        (2, "5.4 实验协议（Experimental Protocol）"),
        (1, "6 仿真实验结果（Simulation Results）"),
        (2, "6.1 位置泛化结果（Position Generalization Results）"),
        (2, "6.2 姿态泛化结果（Orientation Generalization Results）"),
        (2, "6.3 形状泛化结果（Shape Generalization Results）"),
        (2, "6.4 放置任务结果（Placement Results）"),
        (2, "6.5 插入任务结果（Insertion Results）"),
        (2, "6.6 基线比较（Baseline Comparison）"),
        (3, "6.6.1 直接轨迹复现（Naive Replay）"),
        (3, "6.6.2 平移迁移（Translation-only Replay）"),
        (3, "6.6.3 几何感知迁移（Geometry-Aware Replay）"),
        (2, "6.7 消融实验（Ablation Study）"),
        (3, "6.7.1 去除几何对齐（Without Geometric Alignment）"),
        (3, "6.7.2 去除高度映射（Without Height Mapping）"),
        (3, "6.7.3 去除抓取口修正（Without Mouth Anchor Correction）"),
        (1, "7 真实机器人验证（Real Robot Validation）"),
        (2, "7.1 真实机器人系统集成（Real Robot System Integration）"),
        (2, "7.2 真实机器人实验结果（Real Robot Experimental Results）"),
        (2, "7.3 仿真到真实分析（Simulation-to-Real Analysis）"),
        (1, "8 结论（Conclusion）"),
        (2, "8.1 工作总结（Summary）"),
        (2, "8.2 局限与未来工作（Limitations and Future Work）"),
    ]

    anchors = []
    for idx, (level, text) in enumerate(lines, start=1):
        anchor = f"h_{idx:03d}"
        anchors.append((level, text, anchor))
        add_toc_entry(doc, text, anchor, level)

    doc.add_page_break()

    body_title = doc.add_paragraph()
    body_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    body_title.paragraph_format.space_after = Pt(14)
    r = body_title.add_run("正文标题")
    set_font(r, size=15, bold=True, color="000000")

    for bookmark_id, (level, text, anchor) in enumerate(anchors, start=1):
        add_body_heading(doc, text, anchor, bookmark_id, level)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
