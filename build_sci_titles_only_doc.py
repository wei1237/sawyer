from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


OUT = Path(r"D:\ubuntu20\SCI机器人论文目录_仅标题版_20260823.docx")


def set_run_font(run, size=11, bold=False, color="000000", font="微软雅黑"):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run._element.rPr.rFonts.set(qn("w:ascii"), font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para_format(p, left_indent=0, first_line_indent=0, after=3):
    p.paragraph_format.left_indent = Inches(left_indent)
    p.paragraph_format.first_line_indent = Inches(first_line_indent)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.08


def add_line(doc, text, level):
    p = doc.add_paragraph()
    if level == 1:
        set_para_format(p, left_indent=0, after=5)
        run = p.add_run(text)
        set_run_font(run, size=12, bold=True, color="1F4D78")
    elif level == 2:
        set_para_format(p, left_indent=0.28, after=3)
        run = p.add_run(text)
        set_run_font(run, size=10.5, bold=False, color="000000")
    else:
        set_para_format(p, left_indent=0.55, after=2)
        run = p.add_run(text)
        set_run_font(run, size=10, bold=False, color="333333")


def build():
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.9)
    sec.bottom_margin = Inches(0.9)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    style.font.size = Pt(10.5)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("基于少量示教快速迁移的机器人操作泛化方法")
    set_run_font(r, size=16, bold=True, color="0B2545")
    title.paragraph_format.space_after = Pt(10)

    toc_title = doc.add_paragraph()
    toc_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = toc_title.add_run("目录")
    set_run_font(r, size=14, bold=True, color="0B2545")
    toc_title.paragraph_format.space_after = Pt(12)

    lines = [
        (1, "1 引言（Introduction）"),
        (2, "1.1 研究背景与动机（Background and Motivation）"),
        (2, "1.2 现有方法的局限（Limitations of Existing Approaches）"),
        (3, "1.2.1 基于学习的方法（Learning-based Manipulation）"),
        (3, "1.2.2 直接轨迹复现方法（Direct Trajectory Replay）"),
        (3, "1.2.3 几何迁移方法（Geometric Transfer Methods）"),
        (2, "1.3 基于物体相对几何的示范迁移思想（Object-Relative Geometry-Aware Replay）"),
        (2, "1.4 主要贡献（Contributions）"),
        (1, "2 相关工作（Related Work）"),
        (2, "2.1 面向机器人操作的示范学习（Learning from Demonstration for Robot Manipulation）"),
        (2, "2.2 少示范与可泛化机器人操作（Few-Shot and Generalizable Robot Manipulation）"),
        (2, "2.3 几何表示与任务迁移（Geometric Representation and Task Transfer）"),
        (1, "3 问题定义（Problem Formulation）"),
        (2, "3.1 示范表示（Demonstration Representation）"),
        (2, "3.2 当前场景表示（Live Scene Representation）"),
        (2, "3.3 几何迁移目标（Objective of Geometric Transfer）"),
        (1, "4 方法（Method）"),
        (2, "4.1 系统总体框架（System Overview）"),
        (2, "4.2 示范库与任务检索（Demonstration Library and Retrieval）"),
        (2, "4.3 语义感知与三维几何估计（Semantic Perception and 3D Geometry Estimation）"),
        (2, "4.4 物体相对几何对齐（Object-Relative Geometric Alignment）"),
        (2, "4.5 几何感知局部轨迹适应（Geometry-Aware Local Trajectory Adaptation）"),
        (3, "4.5.1 高度感知映射（Height-Aware Mapping）"),
        (3, "4.5.2 抓取口几何锚定修正（Mouth Anchor Correction）"),
        (3, "4.5.3 局部轨迹融合（Tail Blending）"),
        (2, "4.6 机器人执行（Robot Execution）"),
        (1, "5 实验设置（Experimental Setup）"),
        (2, "5.1 仿真与真实机器人平台（Simulation and Real Robot Platforms）"),
        (3, "5.1.1 仿真平台（Simulation Platform）"),
        (3, "5.1.2 真实机器人平台（Real Robot Platform）"),
        (2, "5.2 任务与示范设计（Tasks and Demonstrations）"),
        (3, "5.2.1 位置抓取泛化（Position Grasp Generalization）"),
        (3, "5.2.2 姿态抓取泛化（Orientation Grasp Generalization）"),
        (3, "5.2.3 形状泛化抓取（Shape Generalization Grasp）"),
        (3, "5.2.4 放置任务迁移（Placement Generalization）"),
        (3, "5.2.5 插入任务迁移（Insertion Generalization）"),
        (2, "5.3 评价指标（Evaluation Metrics）"),
        (2, "5.4 实验协议（Experimental Protocol）"),
        (1, "6 仿真实验结果（Simulation Results）"),
        (2, "6.1 位置泛化结果（Position Generalization Results）"),
        (2, "6.2 姿态泛化结果（Orientation Generalization Results）"),
        (2, "6.3 形状泛化结果（Shape Generalization Results）"),
        (2, "6.4 放置任务结果（Placement Results）"),
        (2, "6.5 插入任务结果（Insertion Results）"),
        (2, "6.6 基线比较（Baseline Comparison）"),
        (3, "6.6.1 直接轨迹复现（Naive Replay）"),
        (3, "6.6.2 平移迁移（Translation-only Replay）"),
        (3, "6.6.3 几何感知迁移（Geometry-Aware Replay）"),
        (2, "6.7 消融实验（Ablation Study）"),
        (3, "6.7.1 去除几何对齐（Without Geometric Alignment）"),
        (3, "6.7.2 去除高度映射（Without Height Mapping）"),
        (3, "6.7.3 去除抓取口修正（Without Mouth Anchor Correction）"),
        (1, "7 真实机器人验证（Real Robot Validation）"),
        (2, "7.1 真实机器人系统集成（Real Robot System Integration）"),
        (2, "7.2 真实机器人实验结果（Real Robot Experimental Results）"),
        (2, "7.3 仿真到真实分析（Simulation-to-Real Analysis）"),
        (1, "8 结论（Conclusion）"),
        (2, "8.1 工作总结（Summary）"),
        (2, "8.2 局限与未来工作（Limitations and Future Work）"),
    ]

    for level, text in lines:
        add_line(doc, text, level)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
